import pdfplumber
import fitz
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_to_word(file_bytes: bytes) -> bytes:
    """Convert PDF to Word (.docx) preserving structure."""
    doc = Document()

    # Style heading
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pages_text = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                pages_text.append(text or "")
    except Exception:
        try:
            fdoc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in fdoc:
                pages_text.append(page.get_text())
            fdoc.close()
        except Exception as e:
            raise RuntimeError(f"Extraction failed: {e}")

    for i, text in enumerate(pages_text):
        if i > 0:
            doc.add_page_break()

        heading = doc.add_heading(f"Page {i + 1}", level=2)
        heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        if text:
            lines = text.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Heuristic: short all-caps lines as subheadings
                if len(line) < 80 and line.isupper():
                    p = doc.add_heading(line, level=3)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(4)
        else:
            doc.add_paragraph("[No extractable text on this page]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
